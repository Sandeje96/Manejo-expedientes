import os
import io
from datetime import datetime
from docx import Document
from flask import current_app


def generar_documento_expediente(expediente, plantilla_path=None):
    """
    Genera un documento Word usando una plantilla existente y reemplazando etiquetas.
    
    Args:
        expediente: Objeto Expediente con todos los datos
        plantilla_path: Ruta a la plantilla Word (opcional)
        
    Returns:
        BytesIO: Stream del documento Word generado
    """
    
    # Ruta por defecto de la plantilla
    if plantilla_path is None:
        plantilla_path = os.path.join(os.path.dirname(__file__), 'templates', 'plantilla_expediente.docx')
    
    # Verificar que existe la plantilla
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"No se encontró la plantilla en: {plantilla_path}")
    
    # Abrir la plantilla
    doc = Document(plantilla_path)
    
    # Crear diccionario con todos los datos del expediente
    datos_reemplazo = _crear_diccionario_datos(expediente)
    
    # Reemplazar etiquetas en párrafos
    _reemplazar_en_paragrafos(doc, datos_reemplazo)
    
    # Reemplazar etiquetas en tablas
    _reemplazar_en_tablas(doc, datos_reemplazo)
    
    # Reemplazar etiquetas en encabezados y pies de página
    _reemplazar_en_headers_footers(doc, datos_reemplazo)
    
    # Guardar en memoria
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream

def generar_documento_visado(expediente, plantilla_path=None):
    """
    Genera un documento Word de visado usando la plantilla de visado y reemplazando etiquetas.
    
    Args:
        expediente: Objeto Expediente con todos los datos
        plantilla_path: Ruta a la plantilla Word de visado (opcional)
        
    Returns:
        BytesIO: Stream del documento Word generado
    """
    
    # Ruta por defecto de la plantilla de visado
    if plantilla_path is None:
        plantilla_path = os.path.join(os.path.dirname(__file__), 'templates', 'plantilla_visado.docx')
    
    # Verificar que existe la plantilla
    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"No se encontró la plantilla de visado en: {plantilla_path}")
    
    # Abrir la plantilla
    doc = Document(plantilla_path)
    
    # Crear diccionario con todos los datos del expediente (reutilizamos la función existente)
    datos_reemplazo = _crear_diccionario_datos(expediente)
    
    # Reemplazar etiquetas en párrafos
    _reemplazar_en_paragrafos(doc, datos_reemplazo)
    
    # Reemplazar etiquetas en tablas
    _reemplazar_en_tablas(doc, datos_reemplazo)
    
    # Reemplazar etiquetas en encabezados y pies de página
    _reemplazar_en_headers_footers(doc, datos_reemplazo)
    
    # Guardar en memoria
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream


def _crear_diccionario_datos(expediente):
    """
    Crea un diccionario con todas las etiquetas posibles y sus valores.
    
    Args:
        expediente: Objeto Expediente
        
    Returns:
        dict: Diccionario con etiquetas y valores
    """
    
    def formatear_monto(monto):
        """Formatea un monto en pesos argentinos."""
        if monto:
            return f"$ {monto:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return "$ 0,00"
    
    def formatear_fecha(fecha):
        """Formatea una fecha."""
        if fecha:
            return fecha.strftime('%d/%m/%Y')
        return '-'
    
    def formatear_datetime(fecha_hora):
        """Formatea una fecha y hora."""
        if fecha_hora:
            return fecha_hora.strftime('%d/%m/%Y %H:%M')
        return '-'
    
    def si_no(valor):
        """Convierte booleano a Sí/No."""
        return 'Sí' if valor else 'No'
    
    # Diccionario completo de reemplazos
    datos = {
        # === DATOS BÁSICOS ===
        '<#id>': str(expediente.id),
        '<#nro_expediente_cpim>': expediente.nro_expediente_cpim or 'No asignado',
        '<#fecha>': formatear_fecha(expediente.fecha),
        '<#profesion>': expediente.profesion or '-',
        '<#formato>': expediente.formato or '-',
        '<#nro_copias>': str(expediente.nro_copias) if expediente.nro_copias else '-',
        '<#tipo_trabajo>': expediente.tipo_trabajo or '-',
        
        # === ACTORES ===
        '<#profesional>': expediente.nombre_profesional or '-',
        '<#nombre_profesional>': expediente.nombre_profesional or '-',
        '<#comitente>': expediente.nombre_comitente or '-',
        '<#nombre_comitente>': expediente.nombre_comitente or '-',
        '<#ubicacion>': expediente.ubicacion or '-',
        '<#partida_inmobiliaria>': expediente.partida_inmobiliaria or '-',
        '<#nro_expediente_municipal>': expediente.nro_expediente_municipal or '-',
        
        # === VISADOS ===
        '<#visado_gas>': si_no(expediente.visado_gas),
        '<#visado_salubridad>': si_no(expediente.visado_salubridad),
        '<#visado_electrica>': si_no(expediente.visado_electrica),
        '<#visado_electromecanica>': si_no(expediente.visado_electromecanica),
        
        # === ESTADOS DE PAGO ===
        '<#estado_pago_sellado>': (expediente.estado_pago_sellado or 'pendiente').capitalize(),
        '<#estado_pago_visado>': (expediente.estado_pago_visado or 'pendiente').capitalize(),
        
        # === MONTOS ===
        '<#tasa_sellado>': formatear_monto(expediente.tasa_sellado_monto),
        '<#tasa_sellado_monto>': formatear_monto(expediente.tasa_sellado_monto),
        '<#tasa_visado_electrica>': formatear_monto(expediente.tasa_visado_electrica_monto),
        '<#tasa_visado_electrica_monto>': formatear_monto(expediente.tasa_visado_electrica_monto),
        '<#tasa_visado_salubridad>': formatear_monto(expediente.tasa_visado_salubridad_monto),
        '<#tasa_visado_salubridad_monto>': formatear_monto(expediente.tasa_visado_salubridad_monto),
        '<#tasa_visado_gas>': formatear_monto(expediente.tasa_visado_gas_monto),
        '<#tasa_visado_gas_monto>': formatear_monto(expediente.tasa_visado_gas_monto),
        '<#tasa_visado_electromecanica>': formatear_monto(expediente.tasa_visado_electromecanica_monto),
        '<#tasa_visado_electromecanica_monto>': formatear_monto(expediente.tasa_visado_electromecanica_monto),
        '<#total_visados>': formatear_monto(expediente.total_visados),
        
        # === FECHAS ===
        '<#fecha_salida>': formatear_fecha(expediente.fecha_salida),
        '<#fecha_finalizado>': formatear_datetime(expediente.fecha_finalizado),
        '<#fecha_creacion>': formatear_datetime(expediente.created_at),
        '<#fecha_actualizacion>': formatear_datetime(expediente.updated_at),
        '<#fecha_generacion>': datetime.now().strftime('%d/%m/%Y %H:%M'),
        '<#fecha_hoy>': datetime.now().strftime('%d/%m/%Y'),
        '<#hora_actual>': datetime.now().strftime('%H:%M'),
        
        # === OTROS DATOS ===
        '<#persona_retira>': expediente.persona_retira or '-',
        '<#nro_caja>': str(expediente.nro_caja) if expediente.nro_caja else '-',
        '<#ruta_carpeta>': expediente.ruta_carpeta or '-',
        '<#whatsapp_profesional>': expediente.whatsapp_profesional or '-',
        '<#whatsapp_tramitador>': expediente.whatsapp_tramitador or '-',
        '<#finalizado>': si_no(expediente.finalizado),
        
        # === DATOS GOP (solo si es digital) ===
        '<#gop_numero>': expediente.gop_numero or 'No asignado' if expediente.formato == 'Digital' else 'No aplica (formato papel)',
        '<#gop_estado>': expediente.gop_estado or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#gop_bandeja_actual>': expediente.gop_bandeja_actual or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#gop_usuario_asignado>': expediente.gop_usuario_asignado or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#gop_fecha_entrada>': formatear_fecha(expediente.gop_fecha_entrada) if expediente.formato == 'Digital' else 'No aplica',
        '<#gop_fecha_en_bandeja>': formatear_fecha(expediente.gop_fecha_en_bandeja) if expediente.formato == 'Digital' else 'No aplica',
        '<#gop_ultima_sincronizacion>': formatear_datetime(expediente.gop_ultima_sincronizacion) if expediente.formato == 'Digital' else 'No aplica',
        
        # === BANDEJAS ESPECÍFICAS ===
        '<#bandeja_cpim_nombre>': expediente.bandeja_cpim_nombre or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_cpim_usuario>': expediente.bandeja_cpim_usuario or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_cpim_fecha>': formatear_fecha(expediente.bandeja_cpim_fecha) if expediente.formato == 'Digital' else 'No aplica',
        
        '<#bandeja_imlauer_nombre>': expediente.bandeja_imlauer_nombre or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_imlauer_usuario>': expediente.bandeja_imlauer_usuario or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_imlauer_fecha>': formatear_fecha(expediente.bandeja_imlauer_fecha) if expediente.formato == 'Digital' else 'No aplica',
        
        '<#bandeja_onetto_nombre>': expediente.bandeja_onetto_nombre or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_onetto_usuario>': expediente.bandeja_onetto_usuario or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_onetto_fecha>': formatear_fecha(expediente.bandeja_onetto_fecha) if expediente.formato == 'Digital' else 'No aplica',
        
        '<#bandeja_profesional_nombre>': expediente.bandeja_profesional_nombre or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_profesional_usuario>': expediente.bandeja_profesional_usuario or '-' if expediente.formato == 'Digital' else 'No aplica',
        '<#bandeja_profesional_fecha>': formatear_fecha(expediente.bandeja_profesional_fecha) if expediente.formato == 'Digital' else 'No aplica',
    }
    
    return datos


def _reemplazar_en_paragrafos(doc, datos_reemplazo):
    """
    Reemplaza etiquetas en todos los párrafos del documento.
    
    Args:
        doc: Documento Word
        datos_reemplazo: Diccionario con etiquetas y valores
    """
    
    for paragraph in doc.paragraphs:
        _reemplazar_en_texto(paragraph, datos_reemplazo)


def _reemplazar_en_tablas(doc, datos_reemplazo):
    """
    Reemplaza etiquetas en todas las tablas del documento.
    
    Args:
        doc: Documento Word
        datos_reemplazo: Diccionario con etiquetas y valores
    """
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _reemplazar_en_texto(paragraph, datos_reemplazo)


def _reemplazar_en_headers_footers(doc, datos_reemplazo):
    """
    Reemplaza etiquetas en encabezados y pies de página.
    
    Args:
        doc: Documento Word
        datos_reemplazo: Diccionario con etiquetas y valores
    """
    
    # Encabezados y pies de página
    for section in doc.sections:
        # Encabezado principal
        header = section.header
        for paragraph in header.paragraphs:
            _reemplazar_en_texto(paragraph, datos_reemplazo)
        
        # Pie de página
        footer = section.footer
        for paragraph in footer.paragraphs:
            _reemplazar_en_texto(paragraph, datos_reemplazo)


def _reemplazar_en_texto(paragraph, datos_reemplazo):
    """
    Reemplaza etiquetas en un párrafo específico, manteniendo el formato.
    
    Args:
        paragraph: Párrafo del documento
        datos_reemplazo: Diccionario con etiquetas y valores
    """
    
    # Obtener el texto completo del párrafo
    texto_completo = paragraph.text
    
    # Verificar si hay etiquetas para reemplazar
    hay_reemplazos = any(etiqueta in texto_completo for etiqueta in datos_reemplazo.keys())
    
    if hay_reemplazos:
        # Realizar todos los reemplazos
        for etiqueta, valor in datos_reemplazo.items():
            texto_completo = texto_completo.replace(etiqueta, str(valor))
        
        # Limpiar el párrafo y agregar el nuevo texto
        # Esto mantiene el formato del párrafo pero reemplaza el contenido
        paragraph.clear()
        run = paragraph.add_run(texto_completo)


def listar_etiquetas_disponibles():
    """
    Retorna una lista de todas las etiquetas disponibles para usar en la plantilla.
    Útil para saber qué etiquetas puedes usar en tu documento Word.
    
    Returns:
        list: Lista de etiquetas disponibles
    """
    
    etiquetas = [
        # Datos básicos
        '<#id>', '<#nro_expediente_cpim>', '<#fecha>', '<#profesion>', '<#formato>',
        '<#nro_copias>', '<#tipo_trabajo>',
        
        # Actores
        '<#profesional>', '<#nombre_profesional>', '<#comitente>', '<#nombre_comitente>',
        '<#ubicacion>', '<#partida_inmobiliaria>', '<#nro_expediente_municipal>',
        
        # Visados
        '<#visado_gas>', '<#visado_salubridad>', '<#visado_electrica>', '<#visado_electromecanica>',
        
        # Estados de pago
        '<#estado_pago_sellado>', '<#estado_pago_visado>',
        
        # Montos
        '<#tasa_sellado>', '<#tasa_visado_electrica>', '<#tasa_visado_salubridad>',
        '<#tasa_visado_gas>', '<#tasa_visado_electromecanica>', '<#total_visados>',
        
        # Fechas
        '<#fecha_salida>', '<#fecha_finalizado>', '<#fecha_creacion>', '<#fecha_actualizacion>',
        '<#fecha_generacion>', '<#fecha_hoy>', '<#hora_actual>',
        
        # Otros datos
        '<#persona_retira>', '<#nro_caja>', '<#ruta_carpeta>',
        '<#whatsapp_profesional>', '<#whatsapp_tramitador>', '<#finalizado>',
        
        # GOP (solo digitales)
        '<#gop_numero>', '<#gop_estado>', '<#gop_bandeja_actual>', '<#gop_usuario_asignado>',
        '<#gop_fecha_entrada>', '<#gop_fecha_en_bandeja>', '<#gop_ultima_sincronizacion>',
        
        # Bandejas específicas
        '<#bandeja_cpim_nombre>', '<#bandeja_cpim_usuario>', '<#bandeja_cpim_fecha>',
        '<#bandeja_imlauer_nombre>', '<#bandeja_imlauer_usuario>', '<#bandeja_imlauer_fecha>',
        '<#bandeja_onetto_nombre>', '<#bandeja_onetto_usuario>', '<#bandeja_onetto_fecha>',
        '<#bandeja_profesional_nombre>', '<#bandeja_profesional_usuario>', '<#bandeja_profesional_fecha>',
    ]
    
    return sorted(etiquetas)